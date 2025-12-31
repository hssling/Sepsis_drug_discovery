"""Generate Cover Letter DOCX for IJCCM"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

BASE_DIR = Path(__file__).parent.parent

def create_cover_letter():
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Date
    p = doc.add_paragraph()
    p.add_run('December 31, 2024')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    # To
    doc.add_paragraph('To,')
    doc.add_paragraph('The Editor-in-Chief')
    doc.add_paragraph('Indian Journal of Critical Care Medicine')
    doc.add_paragraph('Official Publication of the Indian Society of Critical Care Medicine')
    
    doc.add_paragraph()
    
    # Subject
    p = doc.add_paragraph()
    p.add_run('Subject: ').bold = True
    p.add_run('Submission of Original Research Manuscript')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Dear Editor,')
    
    doc.add_paragraph()
    
    # Body
    p = doc.add_paragraph()
    p.add_run('We are pleased to submit our manuscript entitled ')
    title_run = p.add_run('Phase-Specific Host-Directed Therapy Targets in Sepsis: An Integrated Multi-omics and Chemoinformatics Pipeline Identifies IL-6, NLRP3, and PD-1 as Priority Candidates')
    title_run.italic = True
    p.add_run(' for consideration for publication in the Indian Journal of Critical Care Medicine.')
    
    doc.add_paragraph()
    
    # Summary section
    p = doc.add_paragraph()
    p.add_run('SUMMARY').bold = True
    
    doc.add_paragraph('Sepsis remains the leading cause of death in intensive care units globally, with approximately 48.9 million cases and 11 million deaths annually. Despite over 100 clinical trials of immunomodulatory agents, no new drugs targeting the host immune response have been approved since the 1960s. This treatment gap represents a critical unmet medical need, particularly in resource-limited settings.')
    
    doc.add_paragraph()
    
    # Key Innovations
    p = doc.add_paragraph()
    p.add_run('KEY INNOVATIONS').bold = True
    
    doc.add_paragraph('1. Phase-Specific Targeting: We address the fundamental reason for past trial failures - the biphasic immune response in sepsis. Our pipeline stratifies targets by immune phase (early hyperinflammation vs. late immunosuppression), enabling precision timing of interventions.')
    
    doc.add_paragraph('2. Immediately Translatable Findings: We identify FDA-approved drugs (Tocilizumab, Baricitinib, Anakinra, Colchicine) with existing safety profiles that can be rapidly repurposed through clinical trials.')
    
    doc.add_paragraph('3. Computational Reproducibility: Our open-source pipeline (https://github.com/hssling/Sepsis_drug_discovery) enables independent validation and extension.')
    
    doc.add_paragraph()
    
    # Relevance
    p = doc.add_paragraph()
    p.add_run('RELEVANCE TO IJCCM').bold = True
    
    doc.add_paragraph('This work is highly relevant to the Indian critical care community:')
    doc.add_paragraph('  - High sepsis burden in India with significant proportion of global cases')
    doc.add_paragraph('  - Cost-effective approach: Drug repurposing reduces development costs')
    doc.add_paragraph('  - Actionable findings: Priority drugs are already available in Indian hospitals')
    
    doc.add_paragraph()
    
    # Declarations
    p = doc.add_paragraph()
    p.add_run('DECLARATIONS').bold = True
    
    doc.add_paragraph('  - This manuscript has not been published previously and is not under consideration elsewhere')
    doc.add_paragraph('  - All data are from publicly available databases (GEO, ChEMBL, Open Targets)')
    doc.add_paragraph('  - No human subjects were involved (computational study)')
    doc.add_paragraph('  - No conflicts of interest to declare')
    doc.add_paragraph('  - The author approves the manuscript for submission')
    
    doc.add_paragraph()
    
    # Closing
    doc.add_paragraph('We believe this work makes a significant contribution to sepsis therapeutics and is well-suited to the scope of IJCCM. We look forward to your favorable consideration.')
    
    doc.add_paragraph()
    doc.add_paragraph('Respectfully submitted,')
    doc.add_paragraph()
    
    # Signature
    p = doc.add_paragraph()
    p.add_run('Dr. Siddalingaiah H S').bold = True
    doc.add_paragraph('Professor, Department of Community Medicine')
    doc.add_paragraph('Shridevi Institute of Medical Sciences and Research Hospital')
    doc.add_paragraph('Tumkur - 572106, Karnataka, India')
    doc.add_paragraph('Email: hssling@yahoo.com')
    doc.add_paragraph('Phone: +91-8941087719')
    doc.add_paragraph('ORCID: 0000-0002-4771-8285')
    
    # Save
    output_path = BASE_DIR / 'manuscripts' / 'CoverLetter_IJCCM.docx'
    doc.save(str(output_path))
    print(f'Created: {output_path}')

if __name__ == '__main__':
    create_cover_letter()
